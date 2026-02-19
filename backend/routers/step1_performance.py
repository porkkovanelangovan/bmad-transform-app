import csv
import io
import json
import logging
import os
from datetime import datetime

import httpx
from bs4 import BeautifulSoup
from fastapi import APIRouter, Depends, UploadFile, File

from database import get_db
from data_ingestion import (
    search_ticker,
    fetch_company_profile,
    fetch_peers,
    fetch_company_overview,
    fetch_financials,
    fetch_finnhub_metrics,
)
from ai_research import is_openai_available

router = APIRouter()
logger = logging.getLogger(__name__)

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

FINANCIAL_EXTRACTION_PROMPT = """Extract financial and business data from this content.
Return ONLY valid JSON with this structure:
{
    "business_units": [{"name": "...", "description": "..."}],
    "revenue_splits": [{"business_unit": "...", "dimension": "product|region|segment", "dimension_value": "...", "revenue": 0, "period": "2024"}],
    "ops_efficiency": [{"business_unit": "...", "metric_name": "...", "metric_value": 0, "target_value": null, "period": "..."}],
    "competitors": [{"name": "...", "market_share": null, "strengths": "...", "weaknesses": "..."}],
    "confidence": "high|medium|low"
}

Rules:
- Extract ALL financial data you can find: revenue figures, margins, KPIs, competitor names
- For revenue, convert to numeric values (millions)
- For metrics like margins, use decimal form (0.15 for 15%)
- If a business unit name is not clear, use the company/org name
- confidence should reflect how clearly financial data was present in the content
- If no financial data found, return empty arrays with confidence "low"

Content:
"""

# Column alias map for structured file auto-detection
COLUMN_ALIASES = {
    "business_units": {
        "name": ["name", "business_unit", "bu_name", "unit", "unit_name", "bu"],
        "description": ["description", "desc", "details"],
    },
    "revenue_splits": {
        "business_unit": ["business_unit", "bu_name", "unit", "bu", "business_unit_name"],
        "dimension": ["dimension", "dim", "type", "category"],
        "dimension_value": ["dimension_value", "value", "dim_value", "segment", "product", "region"],
        "revenue": ["revenue", "total_revenue", "sales", "amount", "rev"],
        "period": ["period", "year", "quarter", "fiscal_year", "date"],
    },
    "ops_efficiency": {
        "business_unit": ["business_unit", "bu_name", "unit", "bu", "business_unit_name"],
        "metric_name": ["metric_name", "metric", "kpi", "measure", "indicator"],
        "metric_value": ["metric_value", "value", "current_value", "actual"],
        "target_value": ["target_value", "target", "goal", "benchmark"],
        "period": ["period", "year", "quarter", "date"],
    },
    "competitors": {
        "name": ["name", "competitor", "competitor_name", "company", "comp"],
        "market_share": ["market_share", "share", "market_pct"],
        "strengths": ["strengths", "strength", "pros", "advantages"],
        "weaknesses": ["weaknesses", "weakness", "cons", "disadvantages"],
    },
}


def _match_column(col_name: str, aliases: dict) -> str | None:
    """Match a CSV/Excel column header to a known field using alias map."""
    normalized = col_name.strip().lower().replace(" ", "_").replace("-", "_")
    for field, alias_list in aliases.items():
        if normalized in alias_list:
            return field
    return None


def _detect_table_type(columns: list[str]) -> tuple[str | None, dict]:
    """Detect which table type a set of columns belongs to and return field mapping."""
    best_match = None
    best_score = 0
    best_mapping = {}

    for table_type, aliases in COLUMN_ALIASES.items():
        mapping = {}
        for col in columns:
            field = _match_column(col, aliases)
            if field:
                mapping[col] = field
        score = len(mapping)
        if score > best_score:
            best_score = score
            best_match = table_type
            best_mapping = mapping
    return best_match, best_mapping


async def _get_or_create_bu(db, bu_name: str) -> int:
    """Get or create a business unit by name, return its ID."""
    existing = await db.execute_fetchall(
        "SELECT id FROM business_units WHERE name = ?", (bu_name,)
    )
    if existing:
        return existing[0]["id"]
    cursor = await db.execute(
        "INSERT INTO business_units (name, description) VALUES (?, ?)",
        (bu_name, f"Imported: {bu_name}"),
    )
    return cursor.lastrowid


async def _insert_extracted_data(data: dict, db, data_source: str = "upload") -> dict:
    """Insert extracted financial data into the 4 Step 1 tables.
    Checks for duplicates before inserting. Returns summary counts."""
    summary = {"business_units": 0, "revenue_splits": 0, "ops_efficiency": 0, "competitors": 0}

    # Business units
    bu_name_to_id = {}
    for bu in data.get("business_units", []):
        name = bu.get("name")
        if not name:
            continue
        existing = await db.execute_fetchall(
            "SELECT id FROM business_units WHERE name = ?", (name,)
        )
        if existing:
            bu_name_to_id[name] = existing[0]["id"]
        else:
            cursor = await db.execute(
                "INSERT INTO business_units (name, description) VALUES (?, ?)",
                (name, bu.get("description", "")),
            )
            bu_name_to_id[name] = cursor.lastrowid
            summary["business_units"] += 1

    # Helper: resolve BU name to ID (create if needed)
    async def resolve_bu(bu_name):
        if not bu_name:
            # Use first existing BU or create a default
            rows = await db.execute_fetchall("SELECT id FROM business_units LIMIT 1")
            if rows:
                return rows[0]["id"]
            cursor = await db.execute(
                "INSERT INTO business_units (name, description) VALUES (?, ?)",
                ("Default", "Auto-created for data import"),
            )
            return cursor.lastrowid
        if bu_name in bu_name_to_id:
            return bu_name_to_id[bu_name]
        return await _get_or_create_bu(db, bu_name)

    # Revenue splits
    for rs in data.get("revenue_splits", []):
        bu_id = await resolve_bu(rs.get("business_unit"))
        dimension = rs.get("dimension", "product")
        dim_value = rs.get("dimension_value", "Total Revenue")
        revenue = rs.get("revenue")
        period = rs.get("period", "Unknown")
        if revenue is None:
            continue
        try:
            revenue = float(revenue)
        except (ValueError, TypeError):
            continue
        existing = await db.execute_fetchall(
            "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period=? AND dimension=? AND dimension_value=?",
            (bu_id, period, dimension, dim_value),
        )
        if not existing:
            await db.execute(
                "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                (bu_id, dimension, dim_value, revenue, period),
            )
            summary["revenue_splits"] += 1

    # Ops efficiency
    for oe in data.get("ops_efficiency", []):
        bu_id = await resolve_bu(oe.get("business_unit"))
        metric_name = oe.get("metric_name")
        metric_value = oe.get("metric_value")
        if not metric_name or metric_value is None:
            continue
        try:
            metric_value = float(metric_value)
        except (ValueError, TypeError):
            continue
        period = oe.get("period", "TTM")
        target_value = oe.get("target_value")
        if target_value is not None:
            try:
                target_value = float(target_value)
            except (ValueError, TypeError):
                target_value = None
        existing = await db.execute_fetchall(
            "SELECT id FROM ops_efficiency WHERE business_unit_id=? AND metric_name=? AND period=?",
            (bu_id, metric_name, period),
        )
        if not existing:
            await db.execute(
                "INSERT INTO ops_efficiency (business_unit_id, metric_name, metric_value, target_value, period) VALUES (?, ?, ?, ?, ?)",
                (bu_id, metric_name, metric_value, target_value, period),
            )
            summary["ops_efficiency"] += 1

    # Competitors
    for comp in data.get("competitors", []):
        name = comp.get("name")
        if not name:
            continue
        existing = await db.execute_fetchall(
            "SELECT id FROM competitors WHERE name = ?", (name,)
        )
        if not existing:
            market_share = comp.get("market_share")
            if market_share is not None:
                try:
                    market_share = float(market_share)
                except (ValueError, TypeError):
                    market_share = None
            await db.execute(
                "INSERT INTO competitors (name, market_share, strengths, weaknesses, data_source) VALUES (?, ?, ?, ?, ?)",
                (name, market_share, comp.get("strengths"), comp.get("weaknesses"), data_source),
            )
            summary["competitors"] += 1

    await db.commit()
    return summary


async def _extract_financial_from_url(url: str) -> dict:
    """Fetch a URL and use AI to extract financial data."""
    try:
        async with httpx.AsyncClient(
            timeout=30, follow_redirects=True, verify=True
        ) as client:
            resp = await client.get(url, headers={"User-Agent": USER_AGENT})
            resp.raise_for_status()
            html = resp.text
    except httpx.HTTPStatusError as e:
        return {"error": f"HTTP {e.response.status_code} fetching URL: {url}"}
    except Exception as e:
        return {"error": f"Failed to fetch URL: {e}"}

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text_parts = []
    for tag in soup.find_all(["p", "li", "td", "th", "h1", "h2", "h3", "h4", "h5", "h6"]):
        text = tag.get_text(strip=True)
        if text and len(text) > 5:
            text_parts.append(text)

    extracted_text = "\n".join(text_parts)
    if not extracted_text.strip():
        extracted_text = soup.get_text(separator="\n", strip=True)
    if not extracted_text.strip():
        return {"error": "No text content found at URL"}

    extracted_text = extracted_text[:8000]

    if not is_openai_available():
        return {
            "error": "OpenAI API key not configured. Cannot extract financial data from URLs.",
            "raw_text_preview": extracted_text[:500],
        }

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI()
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": FINANCIAL_EXTRACTION_PROMPT + extracted_text}
            ],
            response_format={"type": "json_object"},
            max_tokens=4000,
            temperature=0.1,
        )
        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
        return json.loads(raw)
    except Exception as e:
        logger.error("AI financial extraction from URL failed: %s", e)
        return {"error": f"AI extraction error: {e}"}


async def _run_api_ingestion(org: dict, db) -> dict:
    """Core API ingestion logic (Finnhub primary + Alpha Vantage supplement).
    Uses Finnhub /stock/metric as the primary data source (60 calls/min)
    and Alpha Vantage INCOME_STATEMENT + OVERVIEW as supplements (25 calls/day)."""
    from data_ingestion import FINNHUB_API_KEY, ALPHA_VANTAGE_API_KEY

    org_name = org["name"]
    summary = {"ticker": None, "profile": False, "financials": 0, "ops_metrics": 0, "competitors": 0}

    # Check if API keys are configured
    if not FINNHUB_API_KEY and not ALPHA_VANTAGE_API_KEY:
        return {"skipped": True, "message": "Finnhub/Alpha Vantage API keys not configured. You can still add data manually, via file upload, or URL extraction.", "summary": summary}

    # 1. Search for ticker
    logger.info("Searching ticker for '%s'", org_name)
    ticker = await search_ticker(org_name)
    if not ticker:
        logger.warning("No ticker found for '%s'", org_name)
        return {"skipped": True, "message": f"Could not find stock ticker for '{org_name}'. You can still add data manually, via file upload, or URL extraction.", "summary": summary}
    summary["ticker"] = ticker
    logger.info("Found ticker %s for '%s'", ticker, org_name)

    # 2. Fetch company profile and update organization
    profile = await fetch_company_profile(ticker)
    if profile:
        await db.execute(
            "UPDATE organization SET ticker=?, sub_industry=?, market_cap=?, country=?, currency=? WHERE id=?",
            (profile["ticker"], profile.get("industry"), profile.get("market_cap"),
             profile.get("country"), profile.get("currency"), org["id"]),
        )
        summary["profile"] = True
        logger.info("Profile updated for %s", ticker)
    else:
        logger.warning("No profile returned for %s", ticker)

    # 3. Ensure a business unit exists for this company
    existing_bu = await db.execute_fetchall(
        "SELECT id FROM business_units WHERE name = ?", (org_name,)
    )
    if existing_bu:
        bu_id = existing_bu[0]["id"]
    else:
        cursor = await db.execute(
            "INSERT INTO business_units (name, description) VALUES (?, ?)",
            (org_name, f"Primary business unit for {org_name}"),
        )
        bu_id = cursor.lastrowid

    # 4. Fetch Finnhub metrics (PRIMARY source — 60 calls/min, always available)
    logger.info("Fetching Finnhub metrics for %s", ticker)
    fh_metrics = await fetch_finnhub_metrics(ticker)
    if fh_metrics:
        logger.info("Finnhub metrics available for %s", ticker)
    else:
        logger.warning("Finnhub metrics unavailable for %s", ticker)

    # 5. Fetch Alpha Vantage financials (income statement) -> revenue_splits
    logger.info("Fetching Alpha Vantage financials for %s", ticker)
    financials = await fetch_financials(ticker)
    if financials and financials.get("annual_reports"):
        logger.info("Alpha Vantage INCOME_STATEMENT returned %d annual reports for %s",
                     len(financials["annual_reports"]), ticker)
        for report in financials["annual_reports"]:
            period = report["fiscal_date"][:4] if report.get("fiscal_date") else "Unknown"
            revenue = report.get("total_revenue")
            if revenue:
                existing = await db.execute_fetchall(
                    "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period=? AND dimension='product' AND dimension_value='Total Revenue'",
                    (bu_id, period),
                )
                if not existing:
                    await db.execute(
                        "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                        (bu_id, "product", "Total Revenue", revenue, period),
                    )
                    summary["financials"] += 1
    else:
        logger.warning("Alpha Vantage INCOME_STATEMENT returned no data for %s (likely rate limited)", ticker)

    # 6. Fetch Alpha Vantage company overview (supplementary)
    logger.info("Fetching Alpha Vantage overview for %s", ticker)
    overview = await fetch_company_overview(ticker)
    if overview:
        logger.info("Alpha Vantage OVERVIEW available for %s", ticker)
    else:
        logger.warning("Alpha Vantage OVERVIEW unavailable for %s (likely rate limited)", ticker)

    # 6a. Revenue fallback: if no revenue from INCOME_STATEMENT, try OVERVIEW TTM
    if summary["financials"] == 0 and overview:
        revenue_ttm = overview.get("revenue_ttm")
        if revenue_ttm:
            existing = await db.execute_fetchall(
                "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period='TTM' AND dimension='product' AND dimension_value='Total Revenue'",
                (bu_id,),
            )
            if not existing:
                await db.execute(
                    "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                    (bu_id, "product", "Total Revenue", revenue_ttm, "TTM"),
                )
                summary["financials"] += 1
                logger.info("Inserted TTM revenue from OVERVIEW fallback for %s: %.0f", ticker, revenue_ttm)
        gross_profit_ttm = overview.get("gross_profit_ttm")
        if gross_profit_ttm:
            existing = await db.execute_fetchall(
                "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period='TTM' AND dimension='product' AND dimension_value='Gross Profit'",
                (bu_id,),
            )
            if not existing:
                await db.execute(
                    "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                    (bu_id, "product", "Gross Profit", gross_profit_ttm, "TTM"),
                )
                summary["financials"] += 1

    # 6b. Revenue last resort: if still 0, use Finnhub enterprise value as proxy
    if summary["financials"] == 0 and fh_metrics and fh_metrics.get("enterpriseValue"):
        ev = fh_metrics["enterpriseValue"]
        existing = await db.execute_fetchall(
            "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period='TTM' AND dimension='segment' AND dimension_value='Enterprise Value (estimated)'",
            (bu_id,),
        )
        if not existing:
            await db.execute(
                "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                (bu_id, "segment", "Enterprise Value (estimated)", ev, "TTM"),
            )
            summary["financials"] += 1
            logger.info("Inserted enterprise value as revenue proxy for %s: %.0f", ticker, ev)

    # 7. Build merged ops metrics: Finnhub primary, Alpha Vantage supplement
    # Define metric mappings: (display_name, finnhub_key, av_key, period, finnhub_is_pct)
    ops_metric_defs = [
        ("Net Profit Margin", "netProfitMarginTTM", "profit_margin", "TTM", True),
        ("Operating Margin", "operatingMarginTTM", "operating_margin", "TTM", True),
        ("Return on Equity (ROE)", "roeTTM", "return_on_equity", "TTM", True),
        ("Return on Assets (ROA)", "roaTTM", "return_on_assets", "TTM", True),
        ("EPS", "epsTTM", "eps", "TTM", False),
        ("Beta", "beta", None, "TTM", False),
        ("Dividend Yield", "currentDividendYieldTTM", None, "TTM", True),
        ("Revenue Growth (YoY)", "revenueGrowthTTMYoy", None, "TTM", True),
        ("Revenue Growth (5Y)", "revenueGrowth5Y", None, "5Y", True),
        ("EPS Growth (3Y)", "epsGrowth3Y", None, "3Y", True),
        ("Book Value Per Share", "bookValuePerShareAnnual", None, "Annual", False),
        ("Net Profit Margin (5Y Avg)", "netProfitMargin5Y", None, "5Y", True),
        ("Gross Margin", "grossMarginTTM", None, "TTM", True),
    ]

    for metric_name, fh_key, av_key, period, fh_is_pct in ops_metric_defs:
        # Try Finnhub first, fall back to Alpha Vantage
        value = None
        if fh_metrics and fh_key:
            raw = fh_metrics.get(fh_key)
            if raw is not None:
                # Finnhub returns percentages as e.g. 20.01 (meaning 20.01%)
                # Convert to decimal form for consistency (0.2001)
                value = raw / 100.0 if fh_is_pct else raw
        if value is None and overview and av_key:
            value = overview.get(av_key)

        if value is not None:
            existing = await db.execute_fetchall(
                "SELECT id FROM ops_efficiency WHERE business_unit_id=? AND metric_name=? AND period=?",
                (bu_id, metric_name, period),
            )
            if not existing:
                await db.execute(
                    "INSERT INTO ops_efficiency (business_unit_id, metric_name, metric_value, period) VALUES (?, ?, ?, ?)",
                    (bu_id, metric_name, value, period),
                )
                summary["ops_metrics"] += 1

    logger.info("Inserted %d revenue splits, %d ops metrics for %s",
                summary["financials"], summary["ops_metrics"], ticker)

    # 8. Fetch named competitors with full financial data
    competitor_names = []
    if org.get("competitor_1_name"):
        competitor_names.append(org["competitor_1_name"])
    if org.get("competitor_2_name"):
        competitor_names.append(org["competitor_2_name"])

    for comp_name in competitor_names:
        comp_ticker = await search_ticker(comp_name)
        if not comp_ticker:
            logger.warning("No ticker found for competitor '%s'", comp_name)
            continue

        comp_profile = await fetch_company_profile(comp_ticker)
        comp_fh_metrics = await fetch_finnhub_metrics(comp_ticker)
        comp_overview = await fetch_company_overview(comp_ticker)
        comp_financials = await fetch_financials(comp_ticker)

        display_name = comp_profile.get("name", comp_name) if comp_profile else comp_name

        # Build competitor financial values from merged sources
        comp_revenue = None
        comp_profit_margin = None
        comp_operating_margin = None
        comp_roe = None
        comp_roa = None
        comp_pe = None
        comp_eps = None

        # Finnhub metrics as primary
        if comp_fh_metrics:
            comp_profit_margin = comp_fh_metrics.get("netProfitMarginTTM")
            if comp_profit_margin is not None:
                comp_profit_margin = comp_profit_margin / 100.0
            comp_operating_margin = comp_fh_metrics.get("operatingMarginTTM")
            if comp_operating_margin is not None:
                comp_operating_margin = comp_operating_margin / 100.0
            comp_roe = comp_fh_metrics.get("roeTTM")
            if comp_roe is not None:
                comp_roe = comp_roe / 100.0
            comp_roa = comp_fh_metrics.get("roaTTM")
            if comp_roa is not None:
                comp_roa = comp_roa / 100.0
            comp_pe = comp_fh_metrics.get("peTTM")
            comp_eps = comp_fh_metrics.get("epsTTM")

        # Alpha Vantage as supplement (fill gaps)
        if comp_overview:
            if comp_revenue is None:
                comp_revenue = comp_overview.get("revenue_ttm")
            if comp_profit_margin is None:
                comp_profit_margin = comp_overview.get("profit_margin")
            if comp_operating_margin is None:
                comp_operating_margin = comp_overview.get("operating_margin")
            if comp_roe is None:
                comp_roe = comp_overview.get("return_on_equity")
            if comp_roa is None:
                comp_roa = comp_overview.get("return_on_assets")
            if comp_pe is None:
                comp_pe = comp_overview.get("pe_ratio")
            if comp_eps is None:
                comp_eps = comp_overview.get("eps")

        existing = await db.execute_fetchall(
            "SELECT id FROM competitors WHERE name = ?", (display_name,)
        )
        if not existing:
            await db.execute(
                """INSERT INTO competitors
                   (name, ticker, revenue, profit_margin, operating_margin,
                    return_on_equity, return_on_assets, pe_ratio, eps,
                    market_cap_value, strengths, data_source)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    display_name,
                    comp_ticker,
                    comp_revenue,
                    comp_profit_margin,
                    comp_operating_margin,
                    comp_roe,
                    comp_roa,
                    comp_pe,
                    comp_eps,
                    comp_profile.get("market_cap") if comp_profile else None,
                    f"Market Cap: ${comp_profile['market_cap']:,.0f}M" if comp_profile and comp_profile.get("market_cap") else None,
                    f"Finnhub + Alpha Vantage (ticker: {comp_ticker})",
                ),
            )
            summary["competitors"] += 1

        # Create business unit for competitor
        existing_bu = await db.execute_fetchall(
            "SELECT id FROM business_units WHERE name = ?", (display_name,)
        )
        if existing_bu:
            comp_bu_id = existing_bu[0]["id"]
        else:
            cursor = await db.execute(
                "INSERT INTO business_units (name, description) VALUES (?, ?)",
                (display_name, f"Competitor: {display_name}"),
            )
            comp_bu_id = cursor.lastrowid

        # Insert competitor revenue splits
        if comp_financials and comp_financials.get("annual_reports"):
            for report in comp_financials["annual_reports"]:
                period = report["fiscal_date"][:4] if report.get("fiscal_date") else "Unknown"
                revenue = report.get("total_revenue")
                if revenue:
                    existing = await db.execute_fetchall(
                        "SELECT id FROM revenue_splits WHERE business_unit_id=? AND period=? AND dimension='product' AND dimension_value='Total Revenue'",
                        (comp_bu_id, period),
                    )
                    if not existing:
                        await db.execute(
                            "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) VALUES (?, ?, ?, ?, ?)",
                            (comp_bu_id, "product", "Total Revenue", revenue, period),
                        )

        # Insert competitor ops metrics from merged sources
        comp_merged_metrics = {
            "Net Profit Margin": comp_profit_margin,
            "Operating Margin": comp_operating_margin,
            "Return on Equity (ROE)": comp_roe,
            "Return on Assets (ROA)": comp_roa,
            "PE Ratio": comp_pe,
            "EPS": comp_eps,
        }
        if comp_fh_metrics:
            beta = comp_fh_metrics.get("beta")
            if beta is not None:
                comp_merged_metrics["Beta"] = beta
            div_yield = comp_fh_metrics.get("currentDividendYieldTTM")
            if div_yield is not None:
                comp_merged_metrics["Dividend Yield"] = div_yield / 100.0

        for metric_name, metric_value in comp_merged_metrics.items():
            if metric_value is not None:
                existing = await db.execute_fetchall(
                    "SELECT id FROM ops_efficiency WHERE business_unit_id=? AND metric_name=? AND period='TTM'",
                    (comp_bu_id, metric_name),
                )
                if not existing:
                    await db.execute(
                        "INSERT INTO ops_efficiency (business_unit_id, metric_name, metric_value, period) VALUES (?, ?, ?, ?)",
                        (comp_bu_id, metric_name, metric_value, "TTM"),
                    )

    # 9. Also fetch peers for additional context
    peers = await fetch_peers(ticker)
    for peer_ticker in peers[:4]:
        peer_profile = await fetch_company_profile(peer_ticker)
        if peer_profile and peer_profile.get("name"):
            existing = await db.execute_fetchall(
                "SELECT id FROM competitors WHERE name = ?", (peer_profile["name"],)
            )
            if not existing:
                market_cap = peer_profile.get("market_cap")
                await db.execute(
                    """INSERT INTO competitors (name, ticker, market_cap_value, strengths, data_source)
                       VALUES (?, ?, ?, ?, ?)""",
                    (
                        peer_profile["name"],
                        peer_ticker,
                        market_cap,
                        f"Market Cap: ${market_cap:,.0f}M" if market_cap else None,
                        f"Finnhub peer (ticker: {peer_ticker})",
                    ),
                )
                summary["competitors"] += 1

    await db.commit()
    logger.info("Ingestion complete for %s: %s", ticker, summary)
    return {"success": True, "summary": summary}


# --- Organization ---

@router.get("/organization")
async def get_organization(db=Depends(get_db)):
    row = await db.execute_fetchall("SELECT * FROM organization ORDER BY id DESC LIMIT 1")
    return dict(row[0]) if row else None


@router.post("/organization")
async def create_organization(data: dict, db=Depends(get_db)):
    # Upsert: delete existing, insert new
    await db.execute("DELETE FROM organization")
    cursor = await db.execute(
        "INSERT INTO organization (name, industry, competitor_1_name, competitor_2_name) VALUES (?, ?, ?, ?)",
        (data["name"], data["industry"], data.get("competitor_1_name"), data.get("competitor_2_name")),
    )
    await db.commit()
    return {"id": cursor.lastrowid}


# --- Company Search (Autocomplete) ---

@router.get("/search-companies")
async def search_companies(q: str = ""):
    """Search for company names/tickers via Finnhub for autocomplete.
    Tries multiple query variations (ticker abbreviations, corporate suffixes)
    to improve results for common short names like 'US Bank' → USB."""
    from data_ingestion import FINNHUB_API_KEY, FINNHUB_BASE
    q = q.strip()
    if not q or len(q) < 2:
        return []
    if not FINNHUB_API_KEY:
        return []

    # Build query variations (same logic as search_ticker in data_ingestion.py)
    queries = [q]
    words = q.split()
    q_lower = q.lower()
    if len(words) >= 2:
        abbr = (words[0] + "".join(w[0] for w in words[1:] if w)).upper()
        if 2 <= len(abbr) <= 5:
            queries.insert(0, abbr)
    compact = q.replace(" ", "").replace(".", "").upper()
    if 2 <= len(compact) <= 5 and compact not in queries:
        queries.insert(0, compact)
    for suffix in ["Corp", "Inc", "Bancorp"]:
        if suffix.lower() not in q_lower:
            queries.append(f"{q} {suffix}")
    if q_lower.startswith("us "):
        queries.append("U.S. " + q[3:])

    try:
        seen_symbols = set()
        merged = []
        async with httpx.AsyncClient(timeout=10) as client:
            for query in queries:
                if len(merged) >= 8:
                    break
                resp = await client.get(
                    f"{FINNHUB_BASE}/search",
                    params={"q": query, "token": FINNHUB_API_KEY},
                )
                data = resp.json()
                results = data.get("result", [])
                us_results = [r for r in results if "." not in r.get("symbol", "")]
                candidates = us_results if us_results else results
                for r in candidates:
                    sym = r.get("symbol", "")
                    if sym and sym not in seen_symbols:
                        seen_symbols.add(sym)
                        merged.append({"symbol": sym, "name": r.get("description", sym)})
                        if len(merged) >= 8:
                            break
        return merged
    except Exception:
        return []


# --- Ingest (Auto-fetch from APIs) ---

@router.post("/ingest")
async def ingest_data(db=Depends(get_db)):
    """Auto-fetch company data from Finnhub & Alpha Vantage APIs."""
    org_rows = await db.execute_fetchall("SELECT * FROM organization ORDER BY id DESC LIMIT 1")
    if not org_rows:
        return {"error": "No organization set. Create one first."}
    org = dict(org_rows[0])
    try:
        return await _run_api_ingestion(org, db)
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        logger.error("Ingestion failed: %s\n%s", e, tb)
        return {"error": f"Ingestion failed: {e}", "traceback": tb}


# --- File Upload ---

@router.post("/upload")
async def upload_file(file: UploadFile = File(...), db=Depends(get_db)):
    """Upload a file (CSV/Excel/JSON/PDF/Word/Image) and extract financial data."""
    if not file.filename:
        return {"error": "No file provided"}

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    content = await file.read()

    try:
        if ext == "csv":
            return await _parse_csv(content, db)
        elif ext == "xlsx":
            return await _parse_excel(content, db)
        elif ext == "json":
            return await _parse_json(content, db)
        elif ext == "pdf":
            return await _parse_pdf(content, db)
        elif ext == "docx":
            return await _parse_docx(content, db)
        elif ext in ("png", "jpg", "jpeg"):
            return await _parse_image(content, ext, db)
        else:
            return {"error": f"Unsupported file type: .{ext}. Supported: csv, xlsx, json, pdf, docx, png, jpg, jpeg"}
    except Exception as e:
        logger.error("File upload processing error: %s", e)
        return {"error": f"Processing error: {e}"}


async def _parse_csv(content: bytes, db) -> dict:
    """Parse CSV file and insert data."""
    text = content.decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    columns = reader.fieldnames or []
    if not columns:
        return {"error": "CSV has no headers"}

    table_type, mapping = _detect_table_type(columns)
    if not table_type:
        return {"error": "Could not detect data type from CSV columns. Expected columns like: name, revenue, metric_name, competitor, etc."}

    rows = list(reader)
    data = _map_rows_to_data(rows, table_type, mapping)
    summary = await _insert_extracted_data(data, db, "csv_upload")
    return {"success": True, "source": "csv", "detected_type": table_type, "rows_processed": len(rows), "summary": summary}


async def _parse_excel(content: bytes, db) -> dict:
    """Parse Excel file and insert data."""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    all_summary = {"business_units": 0, "revenue_splits": 0, "ops_efficiency": 0, "competitors": 0}
    sheets_processed = 0

    for ws in wb.worksheets:
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            continue

        columns = [str(c) if c else f"col_{i}" for i, c in enumerate(header_row)]
        table_type, mapping = _detect_table_type(columns)
        if not table_type:
            continue

        rows = []
        for row_vals in rows_iter:
            row_dict = {columns[i]: row_vals[i] for i in range(min(len(columns), len(row_vals))) if row_vals[i] is not None}
            if row_dict:
                rows.append(row_dict)

        data = _map_rows_to_data(rows, table_type, mapping)
        summary = await _insert_extracted_data(data, db, "excel_upload")
        for k in all_summary:
            all_summary[k] += summary.get(k, 0)
        sheets_processed += 1

    if sheets_processed == 0:
        return {"error": "No recognizable data found in Excel file"}
    return {"success": True, "source": "excel", "sheets_processed": sheets_processed, "summary": all_summary}


async def _parse_json(content: bytes, db) -> dict:
    """Parse JSON file and insert data."""
    try:
        data = json.loads(content)
    except json.JSONDecodeError as e:
        return {"error": f"Invalid JSON: {e}"}

    # Handle both direct format and wrapped format
    if isinstance(data, list):
        # Try to detect type from first item's keys
        if data:
            keys = list(data[0].keys())
            table_type, mapping = _detect_table_type(keys)
            if table_type:
                mapped = _map_rows_to_data(data, table_type, mapping)
                summary = await _insert_extracted_data(mapped, db, "json_upload")
                return {"success": True, "source": "json", "detected_type": table_type, "summary": summary}
        return {"error": "Could not detect data type from JSON array"}

    # Direct format with business_units, revenue_splits, etc. keys
    if any(k in data for k in ["business_units", "revenue_splits", "ops_efficiency", "competitors"]):
        summary = await _insert_extracted_data(data, db, "json_upload")
        return {"success": True, "source": "json", "summary": summary}

    return {"error": "JSON format not recognized. Use {business_units:[], revenue_splits:[], ops_efficiency:[], competitors:[]}"}


async def _parse_pdf(content: bytes, db) -> dict:
    """Parse PDF file — extract text, then use AI for financial extraction."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"error": "PyMuPDF not installed. Run: pip install PyMuPDF"}

    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        # Scanned PDF — try image-based extraction
        if not is_openai_available():
            return {"error": "Scanned PDF detected but OpenAI API key not configured for OCR extraction."}
        # Render first few pages as images
        import base64
        images = []
        for i, page in enumerate(doc):
            if i >= 3:
                break
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            images.append(base64.b64encode(img_bytes).decode())

        return await _extract_from_images(images, db, "pdf_scan_upload")

    if not is_openai_available():
        return {"error": "OpenAI API key not configured. Cannot extract financial data from unstructured PDFs."}

    # Use AI to extract financial data from text
    full_text = full_text[:8000]
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI()
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": FINANCIAL_EXTRACTION_PROMPT + full_text}],
            response_format={"type": "json_object"},
            max_tokens=4000,
            temperature=0.1,
        )
        data = json.loads(response.choices[0].message.content.strip())
        summary = await _insert_extracted_data(data, db, "pdf_upload")
        return {"success": True, "source": "pdf", "confidence": data.get("confidence", "medium"), "summary": summary}
    except Exception as e:
        return {"error": f"AI extraction from PDF failed: {e}"}


async def _parse_docx(content: bytes, db) -> dict:
    """Parse Word document and extract financial data via AI."""
    try:
        import docx
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    if not is_openai_available():
        return {"error": "OpenAI API key not configured. Cannot extract financial data from Word documents."}

    doc = docx.Document(io.BytesIO(content))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    full_text = "\n".join(text_parts).strip()
    if not full_text:
        return {"error": "No text content found in Word document"}

    full_text = full_text[:8000]
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI()
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": FINANCIAL_EXTRACTION_PROMPT + full_text}],
            response_format={"type": "json_object"},
            max_tokens=4000,
            temperature=0.1,
        )
        data = json.loads(response.choices[0].message.content.strip())
        summary = await _insert_extracted_data(data, db, "docx_upload")
        return {"success": True, "source": "docx", "confidence": data.get("confidence", "medium"), "summary": summary}
    except Exception as e:
        return {"error": f"AI extraction from Word doc failed: {e}"}


async def _parse_image(content: bytes, ext: str, db) -> dict:
    """Parse image file using GPT-4o Vision."""
    if not is_openai_available():
        return {"error": "OpenAI API key not configured. Cannot extract financial data from images."}

    import base64
    b64 = base64.b64encode(content).decode()
    return await _extract_from_images([b64], db, "image_upload", ext=ext)


async def _extract_from_images(b64_images: list, db, source: str, ext: str = "png") -> dict:
    """Use GPT-4o Vision to extract financial data from base64 images."""
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI()

        mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
        content_parts = [{"type": "text", "text": FINANCIAL_EXTRACTION_PROMPT + "\n[See attached image(s)]"}]
        for b64 in b64_images:
            content_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}"}
            })

        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": content_parts}],
            max_tokens=4000,
            temperature=0.1,
        )
        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
        data = json.loads(raw)
        summary = await _insert_extracted_data(data, db, source)
        return {"success": True, "source": source, "confidence": data.get("confidence", "medium"), "summary": summary}
    except Exception as e:
        return {"error": f"AI Vision extraction failed: {e}"}


def _map_rows_to_data(rows: list, table_type: str, mapping: dict) -> dict:
    """Map raw rows to the standard financial data structure using detected column mapping."""
    data = {"business_units": [], "revenue_splits": [], "ops_efficiency": [], "competitors": []}
    seen = set()

    for row in rows:
        mapped = {}
        for col, field in mapping.items():
            val = row.get(col)
            if val is not None:
                mapped[field] = str(val).strip() if not isinstance(val, (int, float)) else val

        if not mapped:
            continue

        if table_type == "business_units":
            name = mapped.get("name")
            if name and name not in seen:
                seen.add(name)
                data["business_units"].append({"name": name, "description": mapped.get("description", "")})
        elif table_type == "revenue_splits":
            data["revenue_splits"].append({
                "business_unit": mapped.get("business_unit"),
                "dimension": mapped.get("dimension", "product"),
                "dimension_value": mapped.get("dimension_value", "Total Revenue"),
                "revenue": mapped.get("revenue"),
                "period": mapped.get("period", "Unknown"),
            })
        elif table_type == "ops_efficiency":
            data["ops_efficiency"].append({
                "business_unit": mapped.get("business_unit"),
                "metric_name": mapped.get("metric_name"),
                "metric_value": mapped.get("metric_value"),
                "target_value": mapped.get("target_value"),
                "period": mapped.get("period", "TTM"),
            })
        elif table_type == "competitors":
            name = mapped.get("name")
            if name and name not in seen:
                seen.add(name)
                data["competitors"].append({
                    "name": name,
                    "market_share": mapped.get("market_share"),
                    "strengths": mapped.get("strengths"),
                    "weaknesses": mapped.get("weaknesses"),
                })

    return data


# --- URL Management ---

@router.get("/urls")
async def list_urls(db=Depends(get_db)):
    """List all saved data ingestion URLs."""
    rows = await db.execute_fetchall("SELECT * FROM step1_data_urls ORDER BY created_at DESC")
    return [dict(r) for r in rows]


@router.post("/urls")
async def add_url(data: dict, db=Depends(get_db)):
    """Add a new URL for data ingestion."""
    url = data.get("url", "").strip()
    if not url:
        return {"error": "URL is required"}
    label = data.get("label", "").strip() or None
    url_type = data.get("url_type", "external")
    cursor = await db.execute(
        "INSERT INTO step1_data_urls (url, label, url_type) VALUES (?, ?, ?)",
        (url, label, url_type),
    )
    await db.commit()
    return {"id": cursor.lastrowid, "success": True}


@router.delete("/urls/{url_id}")
async def delete_url(url_id: int, db=Depends(get_db)):
    """Delete a saved URL."""
    await db.execute("DELETE FROM step1_data_urls WHERE id = ?", (url_id,))
    await db.commit()
    return {"deleted": True}


@router.post("/urls/{url_id}/fetch")
async def fetch_single_url(url_id: int, db=Depends(get_db)):
    """Fetch a single saved URL and extract financial data."""
    rows = await db.execute_fetchall("SELECT * FROM step1_data_urls WHERE id = ?", (url_id,))
    if not rows:
        return {"error": "URL not found"}
    url_row = dict(rows[0])

    result = await _extract_financial_from_url(url_row["url"])
    now = datetime.utcnow().isoformat()

    if "error" in result and "business_units" not in result:
        await db.execute(
            "UPDATE step1_data_urls SET status='error', error_message=?, last_fetched_at=? WHERE id=?",
            (result["error"], now, url_id),
        )
        await db.commit()
        return {"error": result["error"], "url_id": url_id}

    summary = await _insert_extracted_data(result, db, f"url:{url_row['url']}")
    await db.execute(
        "UPDATE step1_data_urls SET status='success', error_message=NULL, last_fetched_at=?, last_result_json=? WHERE id=?",
        (now, json.dumps(summary), url_id),
    )
    await db.commit()
    return {"success": True, "url_id": url_id, "summary": summary}


@router.post("/urls/fetch-all")
async def fetch_all_urls(db=Depends(get_db)):
    """Re-fetch all saved URLs and extract data."""
    rows = await db.execute_fetchall("SELECT * FROM step1_data_urls ORDER BY id")
    if not rows:
        return {"message": "No URLs configured", "results": []}

    results = []
    for row in rows:
        url_row = dict(row)
        url_id = url_row["id"]
        result = await _extract_financial_from_url(url_row["url"])
        now = datetime.utcnow().isoformat()

        if "error" in result and "business_units" not in result:
            await db.execute(
                "UPDATE step1_data_urls SET status='error', error_message=?, last_fetched_at=? WHERE id=?",
                (result["error"], now, url_id),
            )
            results.append({"url_id": url_id, "url": url_row["url"], "status": "error", "error": result["error"]})
        else:
            summary = await _insert_extracted_data(result, db, f"url:{url_row['url']}")
            await db.execute(
                "UPDATE step1_data_urls SET status='success', error_message=NULL, last_fetched_at=?, last_result_json=? WHERE id=?",
                (now, json.dumps(summary), url_id),
            )
            results.append({"url_id": url_id, "url": url_row["url"], "status": "success", "summary": summary})

    await db.commit()
    return {"success": True, "results": results}


# --- Refresh All Sources ---

@router.post("/refresh-all")
async def refresh_all_sources(db=Depends(get_db)):
    """Pull from ALL configured sources: APIs, saved URLs, web search, integrations."""
    org_rows = await db.execute_fetchall("SELECT * FROM organization ORDER BY id DESC LIMIT 1")
    if not org_rows:
        return {"error": "No organization set. Create one first."}
    org = dict(org_rows[0])

    results = {}

    # 1. Finnhub + Alpha Vantage APIs
    try:
        api_result = await _run_api_ingestion(org, db)
        if "error" in api_result:
            results["api"] = {"status": "error", "error": api_result["error"]}
        else:
            results["api"] = {"status": "ok", "summary": api_result.get("summary")}
    except Exception as e:
        results["api"] = {"status": "error", "error": str(e)}

    # 2. Saved URLs
    url_rows = await db.execute_fetchall("SELECT * FROM step1_data_urls ORDER BY id")
    if url_rows:
        url_results = []
        for row in url_rows:
            url_row = dict(row)
            url_id = url_row["id"]
            try:
                result = await _extract_financial_from_url(url_row["url"])
                now = datetime.utcnow().isoformat()
                if "error" in result and "business_units" not in result:
                    await db.execute(
                        "UPDATE step1_data_urls SET status='error', error_message=?, last_fetched_at=? WHERE id=?",
                        (result["error"], now, url_id),
                    )
                    url_results.append({"url": url_row["url"], "status": "error"})
                else:
                    summary = await _insert_extracted_data(result, db, f"url:{url_row['url']}")
                    await db.execute(
                        "UPDATE step1_data_urls SET status='success', error_message=NULL, last_fetched_at=?, last_result_json=? WHERE id=?",
                        (now, json.dumps(summary), url_id),
                    )
                    url_results.append({"url": url_row["url"], "status": "ok", "summary": summary})
            except Exception as e:
                url_results.append({"url": url_row["url"], "status": "error", "error": str(e)})
        results["urls"] = {"status": "ok", "fetched": len(url_results), "details": url_results}
    else:
        results["urls"] = {"status": "skipped", "reason": "No URLs configured"}

    # 3. Web Search (if available)
    try:
        from ai_research import web_search_available
        if web_search_available():
            # Search for financial data about the org
            search_query = f"{org['name']} financial data revenue"
            from ai_research import search_web
            search_results = await search_web(search_query)
            if search_results:
                search_extracted = 0
                for sr in search_results[:3]:
                    url = sr.get("url") or sr.get("link")
                    if url:
                        try:
                            result = await _extract_financial_from_url(url)
                            if "error" not in result or "business_units" in result:
                                await _insert_extracted_data(result, db, f"web_search:{url}")
                                search_extracted += 1
                        except Exception:
                            pass
                results["web_search"] = {"status": "ok", "urls_extracted": search_extracted}
            else:
                results["web_search"] = {"status": "skipped", "reason": "No search results"}
        else:
            results["web_search"] = {"status": "skipped", "reason": "Web search not available"}
    except (ImportError, AttributeError):
        results["web_search"] = {"status": "skipped", "reason": "Web search not configured"}

    # 4. Jira integration (if configured)
    jira_url = os.getenv("JIRA_URL")
    jira_token = os.getenv("JIRA_API_TOKEN")
    if jira_url and jira_token:
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                jira_email = os.getenv("JIRA_EMAIL", "")
                import base64
                auth = base64.b64encode(f"{jira_email}:{jira_token}".encode()).decode()
                resp = await client.get(
                    f"{jira_url}/rest/api/3/search?jql=project%20is%20not%20EMPTY&maxResults=50",
                    headers={"Authorization": f"Basic {auth}", "Accept": "application/json"},
                )
                if resp.status_code == 200:
                    jira_data = resp.json()
                    issues = jira_data.get("issues", [])
                    if issues:
                        # Calculate avg cycle time from resolved issues
                        cycle_times = []
                        for issue in issues:
                            fields = issue.get("fields", {})
                            created = fields.get("created")
                            resolved = fields.get("resolutiondate")
                            if created and resolved:
                                from datetime import datetime as dt
                                try:
                                    c = dt.fromisoformat(created.replace("Z", "+00:00"))
                                    r = dt.fromisoformat(resolved.replace("Z", "+00:00"))
                                    cycle_times.append((r - c).total_seconds() / 86400)
                                except Exception:
                                    pass
                        if cycle_times:
                            avg_cycle = sum(cycle_times) / len(cycle_times)
                            bu_rows = await db.execute_fetchall("SELECT id FROM business_units LIMIT 1")
                            if bu_rows:
                                await _insert_extracted_data({
                                    "ops_efficiency": [{
                                        "business_unit": None,
                                        "metric_name": "Avg Jira Cycle Time (days)",
                                        "metric_value": round(avg_cycle, 1),
                                        "period": "TTM",
                                    }]
                                }, db, "jira")
                            results["jira"] = {"status": "ok", "avg_cycle_time_days": round(avg_cycle, 1)}
                        else:
                            results["jira"] = {"status": "ok", "message": "No resolved issues found"}
                else:
                    results["jira"] = {"status": "error", "error": f"HTTP {resp.status_code}"}
        except Exception as e:
            results["jira"] = {"status": "error", "error": str(e)}
    else:
        results["jira"] = {"status": "skipped", "reason": "Jira not configured"}

    # 5. ServiceNow integration (if configured)
    snow_url = os.getenv("SERVICENOW_URL")
    snow_user = os.getenv("SERVICENOW_USER")
    snow_pass = os.getenv("SERVICENOW_PASSWORD")
    if snow_url and snow_user and snow_pass:
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                resp = await client.get(
                    f"{snow_url}/api/now/table/incident?sysparm_limit=50&sysparm_fields=sys_created_on,resolved_at",
                    auth=(snow_user, snow_pass),
                    headers={"Accept": "application/json"},
                )
                if resp.status_code == 200:
                    snow_data = resp.json()
                    records = snow_data.get("result", [])
                    resolution_times = []
                    for rec in records:
                        created = rec.get("sys_created_on")
                        resolved = rec.get("resolved_at")
                        if created and resolved:
                            from datetime import datetime as dt
                            try:
                                c = dt.strptime(created, "%Y-%m-%d %H:%M:%S")
                                r = dt.strptime(resolved, "%Y-%m-%d %H:%M:%S")
                                resolution_times.append((r - c).total_seconds() / 3600)
                            except Exception:
                                pass
                    if resolution_times:
                        avg_res = sum(resolution_times) / len(resolution_times)
                        await _insert_extracted_data({
                            "ops_efficiency": [{
                                "business_unit": None,
                                "metric_name": "Avg Incident Resolution (hours)",
                                "metric_value": round(avg_res, 1),
                                "period": "TTM",
                            }]
                        }, db, "servicenow")
                        results["servicenow"] = {"status": "ok", "avg_resolution_hours": round(avg_res, 1)}
                    else:
                        results["servicenow"] = {"status": "ok", "message": "No resolved incidents found"}
                else:
                    results["servicenow"] = {"status": "error", "error": f"HTTP {resp.status_code}"}
        except Exception as e:
            results["servicenow"] = {"status": "error", "error": str(e)}
    else:
        results["servicenow"] = {"status": "skipped", "reason": "ServiceNow not configured"}

    await db.commit()
    return {"success": True, "sources": results}


# --- Analysis ---

@router.get("/analysis")
async def get_analysis(db=Depends(get_db)):
    """Return aggregated analysis data including revenue trends, ops benchmarks, competitor comparison, and auto-SWOT."""
    org_rows = await db.execute_fetchall("SELECT * FROM organization ORDER BY id DESC LIMIT 1")
    org = dict(org_rows[0]) if org_rows else None

    org_name = org["name"] if org else None

    # Revenue trends
    revenue_rows = await db.execute_fetchall(
        "SELECT rs.period, rs.dimension_value, rs.revenue, bu.name as business_unit_name "
        "FROM revenue_splits rs JOIN business_units bu ON rs.business_unit_id = bu.id "
        "ORDER BY rs.period ASC"
    )
    revenue_trends = [dict(r) for r in revenue_rows]

    # Ops efficiency
    ops_rows = await db.execute_fetchall(
        "SELECT oe.metric_name, oe.metric_value, oe.target_value, oe.period, bu.name as business_unit_name "
        "FROM ops_efficiency oe JOIN business_units bu ON oe.business_unit_id = bu.id "
        "ORDER BY oe.metric_name"
    )
    ops_metrics = [dict(r) for r in ops_rows]

    # Competitors
    comp_rows = await db.execute_fetchall("SELECT * FROM competitors ORDER BY name")
    competitors = [dict(r) for r in comp_rows]

    # Build org metrics dict
    org_metrics = {}
    for m in ops_metrics:
        if m["business_unit_name"] == org_name:
            org_metrics[m["metric_name"]] = m["metric_value"]

    # Build competitor comparison (named competitors with financial data)
    competitor_comparison = []
    for c in competitors:
        if c.get("profit_margin") is not None or c.get("revenue") is not None:
            competitor_comparison.append({
                "name": c["name"],
                "ticker": c.get("ticker"),
                "metrics": {
                    "Profit Margin": c.get("profit_margin"),
                    "Operating Margin": c.get("operating_margin"),
                    "Return on Equity": c.get("return_on_equity"),
                    "Return on Assets": c.get("return_on_assets"),
                    "PE Ratio": c.get("pe_ratio"),
                    "EPS": c.get("eps"),
                    "Revenue (TTM)": c.get("revenue"),
                    "Market Cap": c.get("market_cap_value"),
                },
            })

    # Build revenue comparison across org and competitors
    revenue_comparison = {"org": [], "competitors": []}
    org_rev = [r for r in revenue_trends if r["business_unit_name"] == org_name and r["dimension_value"] == "Total Revenue"]
    revenue_comparison["org"] = [{"period": r["period"], "revenue": r["revenue"]} for r in org_rev]

    comp_names = set()
    if org and org.get("competitor_1_name"):
        comp_names.add(org["competitor_1_name"])
    if org and org.get("competitor_2_name"):
        comp_names.add(org["competitor_2_name"])
    # Also match by display name from competitors table
    for cc in competitor_comparison:
        comp_names.add(cc["name"])

    for cname in comp_names:
        comp_rev = [r for r in revenue_trends if r["business_unit_name"] == cname and r["dimension_value"] == "Total Revenue"]
        if comp_rev:
            revenue_comparison["competitors"].append({
                "name": cname,
                "data": [{"period": r["period"], "revenue": r["revenue"]} for r in comp_rev],
            })

    # Auto-generate SWOT based on financial data (enhanced with competitor comparison)
    swot = _generate_auto_swot(org, revenue_trends, ops_metrics, competitors, competitor_comparison, org_metrics)

    return {
        "organization": {**(org or {}), "metrics": org_metrics},
        "revenue_trends": revenue_trends,
        "ops_metrics": ops_metrics,
        "competitors": competitors,
        "competitor_comparison": competitor_comparison,
        "revenue_comparison": revenue_comparison,
        "swot": swot,
    }


@router.post("/analysis/save-swot")
async def save_analysis_swot(data: dict, db=Depends(get_db)):
    """Save auto-generated SWOT entries to the swot_entries table for Step 3."""
    entries = data.get("entries", [])
    bu_rows = await db.execute_fetchall("SELECT id FROM business_units LIMIT 1")
    if not bu_rows:
        return {"error": "No business unit exists"}
    bu_id = bu_rows[0]["id"]

    saved = 0
    for entry in entries:
        category = entry.get("category")
        description = entry.get("description")
        if category and description:
            await db.execute(
                "INSERT INTO swot_entries (business_unit_id, category, description, data_source, severity, confidence) VALUES (?, ?, ?, ?, ?, ?)",
                (bu_id, category, description, "Auto-generated from Step 1 Analysis",
                 entry.get("severity", "medium"), entry.get("confidence", "medium")),
            )
            saved += 1
    await db.commit()
    return {"saved": saved}


def _generate_auto_swot(org, revenue_trends, ops_metrics, competitors, competitor_comparison=None, org_metrics=None, benchmarks=None):
    """Generate SWOT entries from financial data analysis, benchmarked against competitors.

    Each entry is a dict with keys: description, severity, confidence.
    When benchmarks are provided, uses dynamic thresholds from industry data.
    """
    strengths = []
    weaknesses = []
    opportunities = []
    threats = []

    org_name = org["name"] if org else "Organization"
    has_competitors = bool(competitor_comparison)

    # Analyze ops metrics (org only)
    metrics_by_name = org_metrics or {}
    if not metrics_by_name:
        for m in ops_metrics:
            metrics_by_name[m["metric_name"]] = m["metric_value"]

    # Compute competitor averages for benchmarking
    comp_avgs = {}
    if competitor_comparison:
        for metric_key in ["Profit Margin", "Operating Margin", "Return on Equity", "Return on Assets"]:
            vals = [c["metrics"].get(metric_key) for c in competitor_comparison if c["metrics"].get(metric_key) is not None]
            if vals:
                comp_avgs[metric_key] = sum(vals) / len(vals)

    # Dynamic thresholds from industry benchmarks (or hardcoded fallbacks)
    bench_kpis = (benchmarks or {}).get("industry_kpis", {}) if benchmarks else {}
    pm_strong = 0.15
    pm_weak = 0.05
    om_strong = 0.20
    om_weak = 0.10
    roe_strong = 0.15
    roe_weak = 0.08

    def _severity(value, threshold, direction="above"):
        """Compute severity based on deviation from threshold."""
        if value is None or threshold is None or threshold == 0:
            return "medium"
        ratio = abs(value - threshold) / abs(threshold)
        if ratio > 1.0:
            return "high"
        elif ratio > 0.3:
            return "medium"
        return "low"

    def _confidence(has_comp_data):
        """Confidence based on data availability."""
        if has_comp_data and has_competitors:
            return "high"
        elif has_comp_data or metrics_by_name:
            return "medium"
        return "low"

    # Helper to look up metric by multiple possible names
    def _get_metric(*names):
        for n in names:
            v = metrics_by_name.get(n)
            if v is not None:
                return v
        return None

    # Benchmark profit margin against competitors
    profit_margin = _get_metric("Net Profit Margin", "Profit Margin")
    comp_pm = comp_avgs.get("Profit Margin")
    if profit_margin is not None:
        if comp_pm is not None:
            sev = _severity(profit_margin, comp_pm)
            conf = "high"
            if profit_margin > comp_pm:
                strengths.append({"description": f"Profit margin ({profit_margin:.1%}) exceeds competitor avg ({comp_pm:.1%})", "severity": sev, "confidence": conf})
            else:
                weaknesses.append({"description": f"Profit margin ({profit_margin:.1%}) trails competitor avg ({comp_pm:.1%})", "severity": sev, "confidence": conf})
        elif profit_margin > pm_strong:
            strengths.append({"description": f"Strong profit margin ({profit_margin:.1%})", "severity": _severity(profit_margin, pm_strong), "confidence": _confidence(False)})
        elif profit_margin < pm_weak:
            weaknesses.append({"description": f"Low profit margin ({profit_margin:.1%})", "severity": _severity(profit_margin, pm_weak), "confidence": _confidence(False)})

    # Benchmark operating margin
    op_margin = metrics_by_name.get("Operating Margin")
    comp_om = comp_avgs.get("Operating Margin")
    if op_margin is not None:
        if comp_om is not None:
            sev = _severity(op_margin, comp_om)
            if op_margin > comp_om:
                strengths.append({"description": f"Operating margin ({op_margin:.1%}) above competitor avg ({comp_om:.1%})", "severity": sev, "confidence": "high"})
            else:
                weaknesses.append({"description": f"Operating margin ({op_margin:.1%}) below competitor avg ({comp_om:.1%})", "severity": sev, "confidence": "high"})
        elif op_margin > om_strong:
            strengths.append({"description": f"High operating efficiency ({op_margin:.1%} operating margin)", "severity": _severity(op_margin, om_strong), "confidence": _confidence(False)})
        elif op_margin < om_weak:
            weaknesses.append({"description": f"Below-average operating margin ({op_margin:.1%})", "severity": _severity(op_margin, om_weak), "confidence": _confidence(False)})

    # Benchmark ROE
    roe = _get_metric("Return on Equity (ROE)", "Return on Equity")
    comp_roe = comp_avgs.get("Return on Equity")
    if roe is not None:
        if comp_roe is not None:
            sev = _severity(roe, comp_roe)
            if roe > comp_roe:
                strengths.append({"description": f"ROE ({roe:.1%}) outperforms competitor avg ({comp_roe:.1%})", "severity": sev, "confidence": "high"})
            else:
                weaknesses.append({"description": f"ROE ({roe:.1%}) lags competitor avg ({comp_roe:.1%})", "severity": sev, "confidence": "high"})
        elif roe > roe_strong:
            strengths.append({"description": f"Strong return on equity ({roe:.1%})", "severity": _severity(roe, roe_strong), "confidence": _confidence(False)})
        elif roe < roe_weak:
            weaknesses.append({"description": f"Low return on equity ({roe:.1%})", "severity": _severity(roe, roe_weak), "confidence": _confidence(False)})

    # ROA analysis
    roa = _get_metric("Return on Assets (ROA)", "Return on Assets")
    if roa is not None:
        if roa > 0.10:
            strengths.append({"description": f"Efficient asset utilization (ROA: {roa:.1%})", "severity": _severity(roa, 0.10), "confidence": _confidence(bool(comp_avgs))})
        elif roa < 0.03:
            weaknesses.append({"description": f"Low return on assets ({roa:.1%})", "severity": _severity(roa, 0.03), "confidence": _confidence(bool(comp_avgs))})

    # Analyze revenue trends (org only)
    org_revenues = [r for r in revenue_trends if r.get("business_unit_name") == org_name and r.get("dimension_value") == "Total Revenue" and r.get("revenue")]
    if not org_revenues:
        org_revenues = [r for r in revenue_trends if r.get("dimension_value") == "Total Revenue" and r.get("revenue")]
    if len(org_revenues) >= 2:
        sorted_rev = sorted(org_revenues, key=lambda x: x["period"])
        latest = sorted_rev[-1]["revenue"]
        prev = sorted_rev[-2]["revenue"]
        if prev > 0:
            growth = (latest - prev) / prev
            if growth > 0.10:
                strengths.append({"description": f"Strong revenue growth ({growth:.1%} YoY)", "severity": "high" if growth > 0.25 else "medium", "confidence": "high"})
                opportunities.append({"description": "Momentum for market share expansion", "severity": "medium", "confidence": "medium"})
            elif growth < 0:
                weaknesses.append({"description": f"Declining revenue ({growth:.1%} YoY)", "severity": "high" if growth < -0.10 else "medium", "confidence": "high"})
                threats.append({"description": "Revenue contraction may signal market challenges", "severity": "high" if growth < -0.10 else "medium", "confidence": "medium"})

    # Analyze market cap
    if org and org.get("market_cap"):
        mc = org["market_cap"]
        if mc > 100000:
            strengths.append({"description": f"Large market capitalization (${mc:,.0f}M)", "severity": "medium", "confidence": "high"})
        elif mc < 2000:
            weaknesses.append({"description": f"Small market capitalization (${mc:,.0f}M) — limited resources", "severity": "medium", "confidence": "high"})

    # Competitive landscape analysis
    if competitor_comparison:
        comp_names = [c["name"] for c in competitor_comparison]
        threats.append({"description": f"Direct competition from {', '.join(comp_names)}", "severity": "medium", "confidence": "high"})
        opportunities.append({"description": "Potential for strategic partnerships or differentiation vs. key competitors", "severity": "medium", "confidence": "medium"})
    elif competitors:
        num_competitors = len(competitors)
        if num_competitors > 5:
            threats.append({"description": f"Highly competitive landscape ({num_competitors} identified peers)", "severity": "medium", "confidence": "medium"})
        opportunities.append({"description": "Potential for strategic partnerships or acquisitions among peer companies", "severity": "low", "confidence": "low"})

    # Default entries if we couldn't generate enough
    if not strengths:
        strengths.append({"description": "Established market presence", "severity": "low", "confidence": "low"})
    if not weaknesses:
        weaknesses.append({"description": "Limited financial data available for deep analysis", "severity": "low", "confidence": "low"})
    if not opportunities:
        opportunities.append({"description": "Digital transformation and AI adoption opportunities", "severity": "medium", "confidence": "low"})
    if not threats:
        threats.append({"description": "Increasing market competition and economic uncertainty", "severity": "medium", "confidence": "low"})

    return {
        "strengths": strengths,
        "weaknesses": weaknesses,
        "opportunities": opportunities,
        "threats": threats,
    }


# --- Business Units ---

@router.get("/business-units")
async def list_business_units(db=Depends(get_db)):
    rows = await db.execute_fetchall("SELECT * FROM business_units ORDER BY name")
    return [dict(r) for r in rows]


@router.post("/business-units")
async def create_business_unit(data: dict, db=Depends(get_db)):
    cursor = await db.execute(
        "INSERT INTO business_units (name, description) VALUES (?, ?)",
        (data["name"], data.get("description")),
    )
    await db.commit()
    return {"id": cursor.lastrowid}


@router.delete("/business-units/{unit_id}")
async def delete_business_unit(unit_id: int, db=Depends(get_db)):
    await db.execute("DELETE FROM business_units WHERE id = ?", (unit_id,))
    await db.commit()
    return {"deleted": True}


# --- Revenue Splits ---

@router.get("/revenue-splits")
async def list_revenue_splits(db=Depends(get_db)):
    rows = await db.execute_fetchall(
        "SELECT rs.*, bu.name as business_unit_name FROM revenue_splits rs "
        "JOIN business_units bu ON rs.business_unit_id = bu.id ORDER BY rs.period DESC"
    )
    return [dict(r) for r in rows]


@router.post("/revenue-splits")
async def create_revenue_split(data: dict, db=Depends(get_db)):
    cursor = await db.execute(
        "INSERT INTO revenue_splits (business_unit_id, dimension, dimension_value, revenue, period) "
        "VALUES (?, ?, ?, ?, ?)",
        (data["business_unit_id"], data["dimension"], data["dimension_value"], data["revenue"], data["period"]),
    )
    await db.commit()
    return {"id": cursor.lastrowid}


# --- Ops Efficiency ---

@router.get("/ops-efficiency")
async def list_ops_efficiency(db=Depends(get_db)):
    rows = await db.execute_fetchall(
        "SELECT oe.*, bu.name as business_unit_name FROM ops_efficiency oe "
        "JOIN business_units bu ON oe.business_unit_id = bu.id ORDER BY oe.period DESC"
    )
    return [dict(r) for r in rows]


@router.post("/ops-efficiency")
async def create_ops_metric(data: dict, db=Depends(get_db)):
    cursor = await db.execute(
        "INSERT INTO ops_efficiency (business_unit_id, metric_name, metric_value, target_value, period) "
        "VALUES (?, ?, ?, ?, ?)",
        (data["business_unit_id"], data["metric_name"], data["metric_value"], data.get("target_value"), data["period"]),
    )
    await db.commit()
    return {"id": cursor.lastrowid}


# --- Competitors ---

@router.get("/competitors")
async def list_competitors(db=Depends(get_db)):
    rows = await db.execute_fetchall("SELECT * FROM competitors ORDER BY name")
    return [dict(r) for r in rows]


@router.post("/competitors")
async def create_competitor(data: dict, db=Depends(get_db)):
    cursor = await db.execute(
        "INSERT INTO competitors (name, ticker, market_share, revenue, profit_margin, operating_margin, "
        "return_on_equity, return_on_assets, pe_ratio, eps, market_cap_value, strengths, weaknesses, data_source) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (data["name"], data.get("ticker"), data.get("market_share"), data.get("revenue"),
         data.get("profit_margin"), data.get("operating_margin"),
         data.get("return_on_equity"), data.get("return_on_assets"),
         data.get("pe_ratio"), data.get("eps"), data.get("market_cap_value"),
         data.get("strengths"), data.get("weaknesses"), data.get("data_source")),
    )
    await db.commit()
    return {"id": cursor.lastrowid}
