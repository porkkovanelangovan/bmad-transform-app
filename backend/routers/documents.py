"""
Documents & RAG Router — Knowledge base management + data mode toggle.
"""

import json
import logging

from fastapi import APIRouter, Depends, UploadFile, File, Form
from database import get_db
from rag_engine import store_document, retrieve_relevant_chunks, build_rag_context, is_live_mode
from ai_research import is_openai_available

router = APIRouter()
logger = logging.getLogger(__name__)


# ─── Data Mode Toggle ────────────────────────────────────────────────────────


@router.get("/data-mode")
async def get_data_mode(db=Depends(get_db)):
    """Get the current organization data mode (demo or live)."""
    row = await db.execute_fetchone("SELECT data_mode FROM organization LIMIT 1")
    mode = (row.get("data_mode") if row else "demo") or "demo"
    return {"data_mode": mode}


@router.post("/data-mode")
async def set_data_mode(data: dict, db=Depends(get_db)):
    """Toggle between demo and live mode."""
    mode = data.get("data_mode", "demo")
    if mode not in ("demo", "live"):
        return {"error": "data_mode must be 'demo' or 'live'"}

    await db.execute("UPDATE organization SET data_mode = ?", [mode])
    await db.commit()
    return {"data_mode": mode, "message": f"Switched to {mode} mode"}


# ─── Document Upload ─────────────────────────────────────────────────────────


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    doc_category: str = Form("general"),
    db=Depends(get_db),
):
    """Upload a document to the organization knowledge base (RAG)."""
    if not file.filename:
        return {"error": "No file provided"}

    # Get org
    org = await db.execute_fetchone("SELECT id FROM organization LIMIT 1")
    if not org:
        return {"error": "No organization set up. Complete Org Setup first."}
    org_id = org["id"]

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    content = await file.read()

    # Extract text based on file type
    text = ""
    try:
        if ext in ("txt", "md", "csv"):
            text = content.decode("utf-8-sig", errors="replace")
        elif ext == "pdf":
            try:
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                text = "\n".join(page.get_text() for page in doc).strip()
            except ImportError:
                return {"error": "PyMuPDF not installed for PDF parsing"}
        elif ext == "docx":
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))
                text = "\n".join(parts)
            except ImportError:
                return {"error": "python-docx not installed for DOCX parsing"}
        elif ext in ("xlsx", "xls"):
            try:
                import openpyxl
                import io
                wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
                parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        vals = [str(c) for c in row if c is not None]
                        if vals:
                            parts.append(" | ".join(vals))
                text = "\n".join(parts)
            except ImportError:
                return {"error": "openpyxl not installed for Excel parsing"}
        elif ext == "json":
            text = content.decode("utf-8-sig", errors="replace")
        else:
            return {"error": f"Unsupported file type: .{ext}"}
    except Exception as e:
        logger.error("Document text extraction failed: %s", e)
        return {"error": f"Failed to extract text: {e}"}

    if not text.strip():
        return {"error": "No text content found in the document"}

    # Store document with embeddings
    try:
        doc_id = await store_document(
            db, org_id, file.filename, ext, text,
            doc_category=doc_category,
            upload_source="manual",
        )
        # Count chunks
        row = await db.execute_fetchone(
            "SELECT COUNT(*) as c FROM document_chunks WHERE document_id = ?", [doc_id]
        )
        chunk_count = row["c"] if row else 0

        return {
            "success": True,
            "document_id": doc_id,
            "filename": file.filename,
            "text_length": len(text),
            "chunks": chunk_count,
            "doc_category": doc_category,
        }
    except Exception as e:
        logger.error("Document storage failed: %s", e)
        return {"error": f"Failed to store document: {e}"}


# ─── Document List / Delete ──────────────────────────────────────────────────


@router.get("/list")
async def list_documents(db=Depends(get_db)):
    """List all documents in the knowledge base."""
    rows = await db.execute_fetchall(
        "SELECT d.id, d.filename, d.file_type, d.doc_category, d.upload_source, "
        "d.step_number, d.created_at, LENGTH(d.content_text) as text_length, "
        "(SELECT COUNT(*) FROM document_chunks WHERE document_id = d.id) as chunks, "
        "(SELECT COUNT(*) FROM document_chunks WHERE document_id = d.id AND embedding_json IS NOT NULL) as embedded_chunks "
        "FROM org_documents d ORDER BY d.created_at DESC"
    )
    return [dict(r) for r in rows]


@router.delete("/{doc_id}")
async def delete_document(doc_id: int, db=Depends(get_db)):
    """Delete a document and its chunks from the knowledge base."""
    await db.execute("DELETE FROM document_chunks WHERE document_id = ?", [doc_id])
    await db.execute("DELETE FROM org_documents WHERE id = ?", [doc_id])
    await db.commit()
    return {"deleted": True}


# ─── Semantic Search ──────────────────────────────────────────────────────────


@router.post("/search")
async def search_documents(data: dict, db=Depends(get_db)):
    """Search the knowledge base using semantic similarity."""
    query = data.get("query", "").strip()
    if not query:
        return {"error": "Please provide a search query"}

    org = await db.execute_fetchone("SELECT id FROM organization LIMIT 1")
    org_id = org["id"] if org else None

    chunks = await retrieve_relevant_chunks(
        db, query, org_id=org_id,
        top_k=data.get("top_k", 10),
        doc_category=data.get("doc_category"),
    )

    results = []
    for c in chunks:
        results.append({
            "chunk_text": c.get("chunk_text", ""),
            "filename": c.get("filename", ""),
            "doc_category": c.get("doc_category", ""),
            "similarity": round(c.get("similarity", 0), 3),
        })
    return {"results": results, "total": len(results)}


# ─── RAG Stats ────────────────────────────────────────────────────────────────


@router.get("/stats")
async def rag_stats(db=Depends(get_db)):
    """Get RAG knowledge base statistics."""
    docs = await db.execute_fetchone("SELECT COUNT(*) as c FROM org_documents")
    chunks = await db.execute_fetchone("SELECT COUNT(*) as c FROM document_chunks")
    embedded = await db.execute_fetchone(
        "SELECT COUNT(*) as c FROM document_chunks WHERE embedding_json IS NOT NULL"
    )
    mode = await db.execute_fetchone("SELECT data_mode FROM organization LIMIT 1")

    return {
        "documents": docs["c"] if docs else 0,
        "chunks": chunks["c"] if chunks else 0,
        "embedded_chunks": embedded["c"] if embedded else 0,
        "data_mode": (mode.get("data_mode") if mode else "demo") or "demo",
        "openai_available": is_openai_available(),
    }
